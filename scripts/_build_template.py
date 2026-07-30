#!/usr/bin/env python3
"""
_build_template.py -- one-off generator for templates/refinement_report_template.docx.

Not part of the report-writing workflow itself (see build_report_docx.py for that).
Run this again only if the template needs to be rebuilt from scratch; for day-to-day
restyling (font, size, color, spacing), open the .docx in Word and edit the named
styles directly (Home tab -> right-click a style -> Modify) -- that's the intended
way to change the look, and does not require touching this script or build_report_docx.py.

Usage:
    python3 _build_template.py
"""

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
import os

BASE_FONT = "Arial"
BASE_SIZE = Pt(10)
HEADING_COLOR = RGBColor(0x1F, 0x3A, 0x5F)

OUT_PATH = os.path.join(os.path.dirname(__file__), "..", "templates", "refinement_report_template.docx")


def build_template():
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = BASE_FONT
    normal.font.size = BASE_SIZE
    rpr = normal.element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), BASE_FONT)

    title = doc.styles["Title"]
    title.font.name = BASE_FONT
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = HEADING_COLOR

    h1 = doc.styles["Heading 1"]
    h1.font.name = BASE_FONT
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = HEADING_COLOR

    h2 = doc.styles["Heading 2"]
    h2.font.name = BASE_FONT
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = HEADING_COLOR

    # Disclaimer style: small italic, based on Normal so it inherits Arial/10pt.
    disclaimer = doc.styles.add_style("TOPilotDisclaimer", WD_STYLE_TYPE.PARAGRAPH)
    disclaimer.base_style = normal
    disclaimer.font.italic = True
    disclaimer.font.size = Pt(8)

    # Caption style for plot links / file listings.
    caption = doc.styles.add_style("TOPilotCaption", WD_STYLE_TYPE.PARAGRAPH)
    caption.base_style = normal
    caption.font.size = Pt(9)

    # Table style: bordered, shaded header row, Arial 10pt body -- built by
    # cloning Word's builtin "Table Grid" so borders come for free, then
    # renaming/restyling it under our own name.
    grid = doc.styles["Table Grid"]
    table_style = doc.styles.add_style("TOPilotTable", WD_STYLE_TYPE.TABLE)
    table_style.base_style = grid
    table_style.font.name = BASE_FONT
    table_style.font.size = BASE_SIZE

    # No visible content is added here -- build_report_docx.py loads this file
    # as its base document and appends all real content (title, sections,
    # tables, disclaimer) on top of it, so anything left in the template body
    # itself would leak into the start of every generated report. The table
    # style's header-row shading is applied per-table by build_report_docx.py
    # (see _set_cell_shading there) since python-docx table styles don't
    # carry per-row shading definitions.
    #
    # A brand-new Document() still contains one empty default paragraph in
    # its body -- strip it so nothing (not even a blank line) precedes the
    # real content build_report_docx.py appends.
    body = doc.element.body
    for p in doc.paragraphs:
        body.remove(p._p)

    doc.save(OUT_PATH)
    print(f"Wrote {os.path.abspath(OUT_PATH)}")


if __name__ == "__main__":
    build_template()
