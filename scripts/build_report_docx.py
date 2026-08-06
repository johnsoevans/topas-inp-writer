#!/usr/bin/env python3
"""
build_report_docx.py -- fill templates/refinement_report_template.docx with a
Rietveld/Pawley/quantitative-analysis report, producing a consistently
formatted .docx instead of (or alongside) the usual markdown report.

This does NOT parse an existing .md report. Build the same content you would
have put in the .md report directly as a `sections` list (see the docstring
of `build_report` below) and call `build_report(...)`. Sections are flexible
(heading + paragraphs and/or a table) so this fits both a structural-Rietveld
report (coordinate table, fit-quality section) and a quantitative-analysis
report (weight-percent table, elemental-composition table) without forcing
either into unused fixed fields.

To change the look (font, size, colors, table borders) edit the named styles
in templates/refinement_report_template.docx directly in Word -- this script
only ever refers to styles by name and does not need to change.

Content conventions (apply these when writing the `sections`/`paragraphs`/
`table` content for any report, not just one example):
  - Use the actual "Å" character for angstroms in lattice parameters, beq/ADP
    values, and bond lengths/distances -- not a bare "A". E.g. "a = 9.1569(3)
    Å", table header "beq (Å²)".
  - State a refined-parameter count as "N parameters refined in total: ..."
    (not "N total:" or "Total: N").

Usage as a library (preferred -- call from a short Python snippet you write
for the specific report):

    from build_report_docx import build_report

    build_report(
        title="LaMnO3 Rietveld refinement -- final report",
        date="2026-07-29",
        sections=[
            {"heading": "Method", "level": 1, "paragraphs": ["..."]},
            {"heading": "Fractional coordinates (final, with esd)", "level": 1,
             "table": {"header": ["Site", "x", "y", "z", "beq (Å²)"],
                       "rows": [["La1", "0.0512(3)", "=x", "0.25", "0.80(2)"]]}},
            {"heading": "Fit quality", "level": 1, "paragraphs": [
                "Final Rietveld: Rwp = 7.36%, Rp = 5.16%, GoF = 2.40."]},
        ],
        plots=[("Final Rietveld fit", "https://claude.ai/code/artifact/...")],
        files=["lamno3.inp -- final refined input file"],
        out_path="lamno3_report.docx",
    )

Usage from the command line (smoke test / regenerating the two example
reports bundled with this skill's test examples):

    python3 build_report_docx.py --demo out.docx
"""

import argparse
import os
import re

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import RGBColor, Inches

DEFAULT_TEMPLATE = os.path.join(os.path.dirname(__file__), "..", "templates",
                                 "refinement_report_template.docx")
VERSION_FILE = os.path.join(os.path.dirname(__file__), "..", "version.txt")
TABLE_HEADER_FILL = "D9E2F3"
HYPERLINK_COLOR = RGBColor(0x05, 0x63, 0xC1)


def _topilot_version():
    """First line of version.txt looks like "v1.0.20 from GitHub" -- return "1.0.20".
    version.txt ships with every install, in this format, so no fallback is needed."""
    with open(VERSION_FILE, "r", encoding="utf-8") as f:
        first_line = f.readline().strip()
    return re.match(r"v?(\d+(?:\.\d+)*)", first_line).group(1)


DEFAULT_DISCLAIMER = (
    f"This refinement was performed with TOPilot ({_topilot_version()}), it is the "
    "user's responsibility to judge its scientific content."
)


def _set_cell_shading(cell, hex_color):
    shd = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color,
    })
    cell._tc.get_or_add_tcPr().append(shd)


def _add_hyperlink(paragraph, url, text):
    """python-docx has no built-in hyperlink API; build the run-inside-a
    field-relationship by hand so plot links are real clickable hyperlinks
    (styled via the run's rPr, not a docx character style, to keep this
    self-contained)."""
    part = paragraph.part
    r_id = part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = paragraph._p.makeelement(qn("w:hyperlink"), {qn("r:id"): r_id})
    run = paragraph._p.makeelement(qn("w:r"), {})
    rpr = paragraph._p.makeelement(qn("w:rPr"), {})
    color = paragraph._p.makeelement(qn("w:color"), {qn("w:val"): "0563C1"})
    underline = paragraph._p.makeelement(qn("w:u"), {qn("w:val"): "single"})
    rpr.append(color)
    rpr.append(underline)
    run.append(rpr)
    text_el = paragraph._p.makeelement(qn("w:t"), {})
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _add_table(doc, header, rows):
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "TOPilotTable"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(header):
        hdr_cells[i].text = str(h)
        _set_cell_shading(hdr_cells[i], TABLE_HEADER_FILL)
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph(style="Normal")  # small gap after every table
    return table


def build_report(title, date, sections, out_path, images=None, plots=None, files=None,
                  disclaimer=DEFAULT_DISCLAIMER,
                  template_path=DEFAULT_TEMPLATE):
    """
    sections: ordered list of dicts, each either:
        {"heading": str, "level": 1|2, "paragraphs": [str, ...]}
        {"heading": str, "level": 1|2, "table": {"header": [...], "rows": [[...], ...]}}
      (a section may include both "paragraphs" and "table"; paragraphs render above the
      table by default -- pass "table_first": True to render the table first, e.g. a
      stats table followed by discussion text)
    images: list of (caption, path) tuples -- path to a PNG/JPG file (e.g. from
      plot_xy_matplotlib.py), embedded directly in a "Figures" section, one per page
      width (6 inches), with its caption as a line below it. Use this for a plot that
      should appear in the document itself, not just linked.
    plots: list of (caption, url) tuples, rendered as an "Interactive Plots" section of hyperlinks --
      for a link to an interactive artifact/page, not an embedded image.
    files: list of plain strings, rendered as a "Files" section bullet list.
    disclaimer: final paragraph, styled distinctly (TOPilotDisclaimer); pass "" to omit.
    """
    doc = Document(template_path)

    doc.add_heading(title, level=0)
    doc.add_paragraph(f"Date: {date}", style="Normal")

    for section in sections:
        doc.add_heading(section["heading"], level=section.get("level", 1))
        table_first = section.get("table_first", False)
        if table_first and "table" in section:
            _add_table(doc, section["table"]["header"], section["table"]["rows"])
        for para in section.get("paragraphs", []):
            doc.add_paragraph(para, style="Normal")
        if not table_first and "table" in section:
            _add_table(doc, section["table"]["header"], section["table"]["rows"])

    if images:
        doc.add_heading("Figures", level=1)
        for caption, path in images:
            doc.add_picture(path, width=Inches(6))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap = doc.add_paragraph(caption, style="TOPilotCaption")
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if plots:
        doc.add_heading("Interactive Plots", level=1)
        for caption, url in plots:
            p = doc.add_paragraph(style="TOPilotCaption")
            _add_hyperlink(p, url, caption)

    if files:
        doc.add_heading("Files", level=1)
        for f in files:
            doc.add_paragraph(f, style="List Bullet")

    if disclaimer:
        doc.add_paragraph(disclaimer, style="TOPilotDisclaimer")

    doc.save(out_path)
    return out_path


def _demo(out_path):
    """Regenerates a docx version of zrw2o8_final_report.md's content, as a
    smoke test that the template/script combination produces a sensible,
    consistently styled document."""
    build_report(
        title="ZrW2O8 Rietveld refinement -- final report",
        date="2026-07-25",
        sections=[
            {"heading": "Method", "level": 1, "paragraphs": [
                "Structure from zrw2o8_rt.cif (space group P2_1 3, a ~ 9.16 Å), converted to "
                "a TOPAS str block with cif_to_str.py, correctly tying the 3-fold-axis sites "
                "(Zr1, W1, W2, O3, O4 all as x=y=z).",
                "Data: 6G17PB01.raw (Bruker, Cu Ka1/Ka2, no monochromator, variable divergence "
                "slits), TCHZ peak shape.",
            ]},
            {"heading": "Fractional coordinates (final, with esd)", "level": 1,
             "table": {"header": ["Site", "x", "y", "z", "beq (Å²)"],
                       "rows": [
                           ["Zr1", "0.00158(20)", "=x", "=x", "0.570(22)"],
                           ["W1", "0.34101(10)", "=x", "=x", "0.870(30)"],
                           ["O1", "0.21041(125)", "0.44440(113)", "0.45403(131)", "1.540(133)"],
                       ]}},
            {"heading": "Fit quality", "level": 1, "paragraphs": [
                "Final Rietveld: Rwp = 7.36%, Rp = 5.16%, GoF = 2.40 (full range, 10-150 "
                "degree 2theta, do_errors on).",
                "Mandatory false-minimum check (TCHZ reset to flat 0.0001 and re-refined): "
                "converged to Rwp = 7.374%, essentially identical -- no false minimum found.",
            ]},
        ],
        plots=[
            ("Final Rietveld fit with esds", "https://claude.ai/code/artifact/19ac541a-a255-48ab-8f6b-0989fea4ee54"),
        ],
        files=["zrw2o8_rietveld.inp -- final refinement (also the TOPAS-GUI launch file)"],
        out_path=out_path,
    )
    print(f"Wrote {os.path.abspath(out_path)}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description=__doc__,
                                      formatter_class=argparse.RawDescriptionHelpFormatter)
    parser.add_argument("out", help="Output .docx path")
    parser.add_argument("--demo", action="store_true",
                         help="Ignore other args and write a smoke-test report to `out`")
    args = parser.parse_args()
    if args.demo:
        _demo(args.out)
    else:
        parser.error("build_report_docx.py is meant to be used as a library (import "
                      "build_report) -- run with --demo <out.docx> for a smoke test.")
